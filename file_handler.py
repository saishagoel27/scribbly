import streamlit as st
import os
import logging
from typing import Dict, Optional, Any
from datetime import datetime
from PIL import Image
import io
import docx
from PyPDF2 import PdfReader

from config import Config

logger = logging.getLogger(__name__)

class FileHandler:
    """Enhanced file handler for managing file uploads and processing"""

    def create_upload_interface(self) -> Optional[Dict[str, Any]]:
        """Enhanced file upload interface with better user guidance"""
        
        # Display supported formats with helpful info
        st.markdown("### 📁 Choose Your Study Material")
        st.markdown("**Supported formats:** PDF, Images (JPG, PNG), Text files, Word documents")
        
        uploaded_file = st.file_uploader(
            "Drag and drop or click to upload",
            type=Config.SUPPORTED_FILE_TYPES,
            help=f"Maximum file size: {Config.MAX_FILE_SIZE_MB}MB"
        )

        if uploaded_file is not None:
            return self._process_uploaded_file(uploaded_file)

        # Show helpful tips when no file is uploaded
        st.info("💡 **Tip:** For best results, upload clear documents with good text visibility")
        
        return None

    def _process_uploaded_file(self, uploaded_file) -> Dict[str, Any]:
        """Enhanced file processing with comprehensive metadata"""
        try:
            # Validate file first
            validation = self.validate_file(uploaded_file)
            if not validation['valid']:
                return {"error": validation['error']}

            # Extract comprehensive metadata
            metadata = self._extract_enhanced_metadata(uploaded_file)
            
            # Prepare file data for processing
            file_data = {
                "file_bytes": uploaded_file.getvalue(),
                "content_type": self._get_content_type(uploaded_file.name),
                "needs_ocr": metadata["file_extension"] in ['pdf', 'jpg', 'jpeg', 'png']
            }

            logger.info(f"Processed file: {metadata['filename']} ({metadata['file_size_mb']:.2f}MB)")

            return {
                "file_data": file_data,
                "metadata": metadata,
                "status": "ready_for_processing"
            }

        except Exception as e:
            logger.error(f"File processing error: {e}")
            return {"error": f"File processing failed: {str(e)}"}

    def _extract_enhanced_metadata(self, uploaded_file) -> Dict[str, Any]:
        """Extract comprehensive metadata from uploaded file"""
        
        file_extension = uploaded_file.name.split('.')[-1].lower()
        file_size_mb = uploaded_file.size / (1024 * 1024)
        
        # Base metadata
        metadata = {
            "filename": uploaded_file.name,
            "file_extension": file_extension,
            "file_size_mb": file_size_mb,
            "file_size_bytes": uploaded_file.size,
            "upload_timestamp": datetime.now().isoformat(),
            "estimated_pages": 1,  # Default
            "estimated_reading_time": "1-2 minutes",  # Default
            "content_preview": "",
            "processing_complexity": "low"
        }
        
        try:
            # File-specific metadata extraction
            if file_extension == 'pdf':
                metadata.update(self._extract_pdf_metadata(uploaded_file))
            elif file_extension in ['jpg', 'jpeg', 'png']:
                metadata.update(self._extract_image_metadata(uploaded_file))
            elif file_extension == 'docx':
                metadata.update(self._extract_docx_metadata(uploaded_file))
            elif file_extension == 'txt':
                metadata.update(self._extract_text_metadata(uploaded_file))
                
        except Exception as e:
            logger.warning(f"Metadata extraction failed for {uploaded_file.name}: {e}")
            # Continue with basic metadata if detailed extraction fails
            
        return metadata
    
    def _extract_pdf_metadata(self, uploaded_file) -> Dict[str, Any]:
        """Extract metadata from PDF files"""
        try:
            pdf_reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
            num_pages = len(pdf_reader.pages)
            
            # Estimate reading time (assuming 200 words per page, 200 WPM reading speed)
            estimated_words = num_pages * 200
            reading_time_minutes = max(1, estimated_words // 200)
            
            # Try to extract some preview text
            preview_text = ""
            if num_pages > 0:
                try:
                    first_page = pdf_reader.pages[0]
                    preview_text = first_page.extract_text()[:200] + "..." if first_page.extract_text() else ""
                except:
                    preview_text = "Preview not available"
            
            return {
                "estimated_pages": num_pages,
                "estimated_reading_time": f"{reading_time_minutes}-{reading_time_minutes + 2} minutes",
                "content_preview": preview_text,
                "processing_complexity": "high" if num_pages > 10 else "medium"
            }
            
        except Exception as e:
            logger.warning(f"PDF metadata extraction failed: {e}")
            return {
                "estimated_pages": "Unknown",
                "processing_complexity": "medium"
            }
    
    def _extract_image_metadata(self, uploaded_file) -> Dict[str, Any]:
        """Extract metadata from image files"""
        try:
            image = Image.open(io.BytesIO(uploaded_file.getvalue()))
            width, height = image.size
            
            # Estimate complexity based on image size
            total_pixels = width * height
            if total_pixels > 2000000:  # > 2MP
                complexity = "high"
                reading_time = "2-4 minutes"
            elif total_pixels > 500000:  # > 0.5MP
                complexity = "medium" 
                reading_time = "1-3 minutes"
            else:
                complexity = "low"
                reading_time = "1-2 minutes"
            
            return {
                "image_dimensions": f"{width} x {height}",
                "estimated_reading_time": reading_time,
                "processing_complexity": complexity,
                "content_preview": f"Image file ({width}x{height} pixels)"
            }
            
        except Exception as e:
            logger.warning(f"Image metadata extraction failed: {e}")
            return {"processing_complexity": "medium"}
    
    def _extract_docx_metadata(self, uploaded_file) -> Dict[str, Any]:
        """Extract metadata from Word documents"""
        try:
            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
            
            # Count paragraphs and estimate pages
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            estimated_pages = max(1, paragraph_count // 15)  # Roughly 15 paragraphs per page
            
            # Estimate reading time
            total_text = " ".join([p.text for p in doc.paragraphs])
            word_count = len(total_text.split())
            reading_time_minutes = max(1, word_count // 200)  # 200 WPM
            
            # Get preview text
            preview_text = ""
            for p in doc.paragraphs[:3]:  # First 3 paragraphs
                if p.text.strip():
                    preview_text += p.text[:100] + " "
            preview_text = preview_text[:200] + "..." if len(preview_text) > 200 else preview_text
            
            return {
                "estimated_pages": estimated_pages,
                "estimated_reading_time": f"{reading_time_minutes}-{reading_time_minutes + 1} minutes",
                "content_preview": preview_text or "Document preview not available",
                "word_count": word_count,
                "processing_complexity": "high" if word_count > 2000 else "medium"
            }
            
        except Exception as e:
            logger.warning(f"DOCX metadata extraction failed: {e}")
            return {"processing_complexity": "medium"}
    
    def _extract_text_metadata(self, uploaded_file) -> Dict[str, Any]:
        """Extract metadata from text files"""
        try:
            text_content = uploaded_file.getvalue().decode('utf-8')
            
            word_count = len(text_content.split())
            line_count = len(text_content.split('\n'))
            
            # Estimate pages (assuming 250 words per page)
            estimated_pages = max(1, word_count // 250)
            
            # Estimate reading time
            reading_time_minutes = max(1, word_count // 200)  # 200 WPM
            
            # Get preview
            preview_text = text_content[:200] + "..." if len(text_content) > 200 else text_content
            
            return {
                "estimated_pages": estimated_pages,
                "estimated_reading_time": f"{reading_time_minutes} minutes",
                "content_preview": preview_text,
                "word_count": word_count,
                "line_count": line_count,
                "processing_complexity": "low" if word_count < 500 else "medium"
            }
            
        except UnicodeDecodeError:
            return {
                "content_preview": "Text encoding issue - file may not be properly formatted",
                "processing_complexity": "high"
            }
        except Exception as e:
            logger.warning(f"Text metadata extraction failed: {e}")
            return {"processing_complexity": "medium"}

    def validate_file(self, uploaded_file) -> Dict[str, Any]:
        """Comprehensive file validation with helpful error messages"""
        try:
            # Check file size first
            size_check = Config.validate_file_size(uploaded_file.size)
            if not size_check['valid']:
                return {
                    'valid': False,
                    'error': size_check['error'],
                    'suggestion': "Try compressing your file or splitting large documents into smaller sections."
                }

            # Check file extension
            extension = uploaded_file.name.split('.')[-1].lower()
            if not Config.validate_file_type(extension):
                return {
                    'valid': False,
                    'error': f"Unsupported file type '.{extension}'",
                    'suggestion': f"Please use one of these formats: {', '.join(Config.SUPPORTED_FILE_TYPES).upper()}"
                }

            # Content validation based on file type
            validation_result = self._validate_file_content(uploaded_file, extension)
            if not validation_result['valid']:
                return validation_result

            # Check for empty or very small files
            if uploaded_file.size < 100:  # Less than 100 bytes
                return {
                    'valid': False,
                    'error': "File appears to be empty or too small",
                    'suggestion': "Please upload a file with actual content to process."
                }

            return {'valid': True, 'message': 'File validation passed'}

        except Exception as e:
            logger.error(f"File validation error: {e}")
            return {
                'valid': False,
                'error': f"File validation failed: {str(e)}",
                'suggestion': "Please try uploading a different file."
            }
    
    def _validate_file_content(self, uploaded_file, extension: str) -> Dict[str, Any]:
        """Validate file content integrity"""
        try:
            file_bytes = uploaded_file.getvalue()
            
            if extension in ['jpg', 'jpeg', 'png']:
                try:
                    image = Image.open(io.BytesIO(file_bytes))
                    # Check if image is very small (might not have readable text)
                    if image.size[0] < 100 or image.size[1] < 100:
                        return {
                            'valid': False,
                            'error': "Image resolution too low for text extraction",
                            'suggestion': "Please upload a higher resolution image (at least 100x100 pixels)."
                        }
                except Exception:
                    return {
                        'valid': False,
                        'error': "Invalid or corrupted image file",
                        'suggestion': "Please upload a valid JPG or PNG image."
                    }
            
            elif extension == 'pdf':
                try:
                    pdf_reader = PdfReader(io.BytesIO(file_bytes))
                    if len(pdf_reader.pages) == 0:
                        return {
                            'valid': False,
                            'error': "PDF file has no pages",
                            'suggestion': "Please upload a PDF with content."
                        }
                except Exception:
                    return {
                        'valid': False,
                        'error': "Invalid or corrupted PDF file",
                        'suggestion': "Please upload a valid PDF document."
                    }
            
            elif extension == 'docx':
                try:
                    doc = docx.Document(io.BytesIO(file_bytes))
                    # Check if document has any text content
                    has_text = any(p.text.strip() for p in doc.paragraphs)
                    if not has_text:
                        return {
                            'valid': False,
                            'error': "Word document appears to be empty",
                            'suggestion': "Please upload a document with text content."
                        }
                except Exception:
                    return {
                        'valid': False,
                        'error': "Invalid or corrupted Word document",
                        'suggestion': "Please upload a valid .docx file."
                    }
            
            elif extension == 'txt':
                try:
                    text_content = file_bytes.decode('utf-8')
                    if len(text_content.strip()) < 10:
                        return {
                            'valid': False,
                            'error': "Text file has insufficient content",
                            'suggestion': "Please upload a text file with at least a few sentences."
                        }
                except UnicodeDecodeError:
                    return {
                        'valid': False,
                        'error': "Text file encoding not supported",
                        'suggestion': "Please save your text file as UTF-8 encoded."
                    }
            
            return {'valid': True}
            
        except Exception as e:
            logger.error(f"Content validation error: {e}")
            return {
                'valid': False,
                'error': "File content validation failed",
                'suggestion': "Please try a different file."
            }

    def _get_content_type(self, filename: str) -> str:
        """Get content type for Azure processing"""
        extension = filename.split('.')[-1].lower()
        content_types = {
            'pdf': 'application/pdf',
            'jpg': 'image/jpeg',
            'jpeg': 'image/jpeg',
            'png': 'image/png',
            'txt': 'text/plain',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        return content_types.get(extension, 'application/octet-stream')

# Global Instance
file_handler = FileHandler()